from docx import Document
import logging
import os
from typing import Dict, Any, List, Optional
from app.services.parsers.base_parser import BaseParser

logger = logging.getLogger(__name__)


class DocxParser(BaseParser):
    """Parser especializado en documentos de Word (.docx)."""

    def parse(self, file_path: str) -> Dict[str, Any]:
        """Extrae texto y tablas del documento Word."""
        try:
            if not os.path.exists(file_path):
                return {"error": "File not found"}

            doc = Document(file_path)
            
            # Extraer texto de párrafos
            full_text: List[str] = []
            for para in doc.paragraphs:
                if para.text.strip():
                    full_text.append(para.text.strip())
            
            # Extraer datos de tablas
            tables_data: List[List[List[str]]] = []
            for table in doc.tables:
                table_rows: List[List[str]] = []
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells]
                    if any(row_data):  # Evitar filas vacías
                        table_rows.append(row_data)
                if table_rows:
                    tables_data.append(table_rows)

            return {
                "format": "DOCX",
                "paragraph_count": len(doc.paragraphs),
                "table_count": len(doc.tables),
                "content_text": "\n".join(full_text),
                "tables": tables_data,
                "metadata": self.get_metadata(file_path)
            }
        except Exception as e:
            logger.error(f"Error parseando DOCX {file_path}: {e}")
            return {"error": str(e)}

    def get_metadata(self, file_path: str) -> Dict[str, Any]:
        """Extrae metadatos del archivo Word."""
        try:
            doc = Document(file_path)
            prop = doc.core_properties
            return {
                "author": prop.author,
                "category": prop.category,
                "comments": prop.comments,
                "content_status": prop.content_status,
                "created": prop.created.isoformat() if prop.created else None,
                "identifier": prop.identifier,
                "keywords": prop.keywords,
                "last_modified_by": prop.last_modified_by,
                "language": prop.language,
                "modified": prop.modified.isoformat() if prop.modified else None,
                "subject": prop.subject,
                "title": prop.title,
                "version": prop.version
            }
        except Exception:
            return {}
